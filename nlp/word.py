import docx

def textWord(file):
	doc = docx.Document(file)
	text = []
	for para in doc.paragraphs:
		text.append(para)
	return "\n".join(text)

file = textWord(file)

print('Paragraph 1:',doc.paragraphs[0].text) 
print('Number of runs in paragraph 1:',len(doc.paragraphs[0].runs)) 
for idx, run in enumerate(doc.paragraphs[0].runs):  
	print('Run %s : %s' %(idx,run.text))
	print('is Run 0 underlined:',doc.paragraphs[0].runs[5].underline) 
	print('is Run 2 bold:',doc.paragraphs[0].runs[1].bold) 
	print('is Run 7 italic:',doc.paragraphs[0].runs[3].italic)